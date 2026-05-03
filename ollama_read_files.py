from ollama import chat
import argparse
import os
import pdfplumber
import docx

parser = argparse.ArgumentParser()
parser.add_argument("--target", help="Specific information you're looking for in the file(s)")
parser.add_argument("filelist", help="List of files to read", nargs="+")
parser.add_argument("--model", default="qwen3.6-128k", help="Specify the LLM to use (default: qwen3.5:9B-128k)")
args = parser.parse_args()

abs_path_list = [os.path.abspath(f) for f in args.filelist]

if args.model:
	model = args.model

if args.target:
	content = "Read these file(s) searching for this information: '%s'. Then produce a summary of that information available within each file: %s" % (args.target, ", ".join(abs_path_list))
else:
	content = "Read these file(s) and produce a summary of each: %s" % ", ".join(abs_path_list)


def _read_pdf(path: str) -> list:
	lines = []
	with pdfplumber.open(path) as pdf:
		for i, page in enumerate(pdf.pages, 1):
			text = page.extract_text() or ""
			lines.append("--- Page %d ---" % i)
			lines.extend(text.splitlines())
	return lines


def _read_docx(path: str) -> list:
	doc = docx.Document(path)
	lines = []
	for para in doc.paragraphs:
		lines.append(para.text)
	for table in doc.tables:
		for row in table.rows:
			lines.append("\t".join(cell.text for cell in row.cells))
	return lines


def _read_text(path: str) -> list:
	with open(path, 'r', errors='replace') as fh:
		return fh.read().splitlines(keepends=True)


def get_file_info(path: str) -> str:
	"""Get metadata about a file without reading its contents.
	Always call this before read_file or search_file on an unknown file.

	Args:
		path: Absolute path to the file.

	Returns:
		str: File size in bytes, line count, and format. Use this to decide
			whether to read the whole file or use search_file first.
	"""
	try:
		size = os.path.getsize(path)
		ext = os.path.splitext(path)[1].lower()
		if ext == '.pdf':
			with pdfplumber.open(path) as pdf:
				line_count = sum(
					len((page.extract_text() or "").splitlines()) for page in pdf.pages
				)
				page_count = len(pdf.pages)
			return "path: %s | format: PDF | pages: %d | lines: ~%d | size: %d bytes" % (
				path, page_count, line_count, size)
		elif ext == '.docx':
			doc = docx.Document(path)
			line_count = len(doc.paragraphs)
			return "path: %s | format: DOCX | paragraphs: %d | size: %d bytes" % (
				path, line_count, size)
		else:
			with open(path, 'rb') as fh:
				if b'\x00' in fh.read(1024):
					return "path: %s | format: binary | size: %d bytes (cannot be read as text)" % (path, size)
			with open(path, 'r', errors='replace') as fh:
				line_count = sum(1 for _ in fh)
			return "path: %s | format: text | lines: %d | size: %d bytes" % (path, line_count, size)
	except FileNotFoundError:
		return "Error: file not found: %s" % path
	except Exception as e:
		return "Error getting info for %s: %s" % (path, e)


def search_file(path: str, pattern: str, context_lines: int = 2) -> str:
	"""Search a file for lines matching a regex pattern, like grep -n -C.
	Use this instead of read_file when looking for specific content in a large file.

	Args:
		path: Absolute path to the file to search.
		pattern: Regular expression to search for.
		context_lines: Number of lines to show before and after each match (default 2).

	Returns:
		str: Matching lines with line numbers and surrounding context, separated
			by "--" between non-adjacent matches. Returns a message if no matches found.
	"""
	import re
	try:
		ext = os.path.splitext(path)[1].lower()
		if ext == '.pdf':
			lines = _read_pdf(path)
		elif ext == '.docx':
			lines = _read_docx(path)
		else:
			with open(path, 'rb') as fh:
				if b'\x00' in fh.read(1024):
					return "Error: %s appears to be a binary file." % path
			lines = [l.rstrip('\n') for l in _read_text(path)]

		compiled = re.compile(pattern, re.IGNORECASE)
		match_indices = [i for i, line in enumerate(lines) if compiled.search(line)]

		if not match_indices:
			return "No matches for '%s' in %s" % (pattern, path)

		output = []
		prev_end = -1
		for idx in match_indices:
			start = max(0, idx - context_lines)
			end = min(len(lines), idx + context_lines + 1)
			if start > prev_end + 1:
				output.append("--")
			for i in range(max(start, prev_end + 1), end):
				marker = ">" if i == idx else " "
				output.append("%d%s\t%s" % (i + 1, marker, lines[i]))
			prev_end = end - 1

		return "\n".join(output)
	except FileNotFoundError:
		return "Error: file not found: %s" % path
	except re.error as e:
		return "Invalid regex pattern '%s': %s" % (pattern, e)
	except Exception as e:
		return "Error searching %s: %s" % (path, e)


def read_file(path: str, offset: int = 1, limit: int = 2000) -> str:
	"""Read the contents of a file from the filesystem, with line numbers.
	Supports plain text, LaTeX (.tex), PDF (.pdf), and Word documents (.docx).

	Args:
		path: Absolute path to the file to read.
		offset: Line number to start reading from (1-indexed, default 1).
		limit: Maximum number of lines to return (default 2000). Call again
			with a higher offset to read further into the file.

	Returns:
		str: File contents with each line prefixed by its line number.
			PDF pages are separated by "--- Page N ---" markers.
			Returns an error message if the file cannot be read.
	"""
	try:
		ext = os.path.splitext(path)[1].lower()
		if ext == '.pdf':
			lines = _read_pdf(path)
		elif ext == '.docx':
			lines = _read_docx(path)
		else:
			with open(path, 'rb') as fh:
				if b'\x00' in fh.read(1024):
					return "Error: %s appears to be a binary file." % path
			lines = _read_text(path)

		start = max(0, offset - 1)
		chunk = lines[start:start + limit]
		if not chunk:
			return "Error: offset %d is beyond end of file (%d lines)." % (offset, len(lines))
		return "".join("%d\t%s\n" % (start + i + 1, line.rstrip('\n')) for i, line in enumerate(chunk))
	except FileNotFoundError:
		return "Error: file not found: %s" % path
	except Exception as e:
		return "Error reading %s: %s" % (path, e)


messages = [{'role': 'user', 'content': content}]

# Agentic loop: keep going until the model stops calling tools
while True:
	response = chat(model=model, messages=messages, tools=[get_file_info, search_file, read_file])
	messages.append(response.message)

	if not response.message.tool_calls:
		break

	for tool_call in response.message.tool_calls:
		fn_name = tool_call.function.name
		fn_args = tool_call.function.arguments
		print("Tool call: %s(%s)" % (fn_name, fn_args))

		tools = {'read_file': read_file, 'search_file': search_file, 'get_file_info': get_file_info}
		if fn_name in tools:
			try:
				result = tools[fn_name](**fn_args)
			except Exception as e:
				result = "Error in %s: %s" % (fn_name, e)
		else:
			result = "Unknown tool: %s" % fn_name

		messages.append({'role': 'tool', 'content': result})

print(response.message.content)
